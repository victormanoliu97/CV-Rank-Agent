"""Utility functions to load and extract text from CV files (PDF and DOCX)."""

def load_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    import pymupdf

    doc = pymupdf.open(file_path)
    parts: list[str] = []

    for page in doc:
        text = page.get_text().strip()
        if text:
            parts.append(text)

    doc.close()
    return "\n".join(parts)

def load_docx(file_path: str) -> str:
    """Extract text from a DOCX file.

    Extracts content from paragraphs, tables, headers, and footers
    to capture as much CV text as possible.
    """
    from docx import Document

    doc = Document(file_path)
    parts: list[str] = []

    # --- header / footer text (each section can have its own) ---
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer and header_footer.is_linked_to_previous is False:
                for para in header_footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    # --- body paragraphs ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # --- tables (row by row, tab-separated cells) ---
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)

def load_cv(file_path: str) -> str:
    """Load CV text from a PDF or DOCX file."""
    if file_path.lower().endswith(".pdf"):
        return load_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported CV file format: {file_path}")